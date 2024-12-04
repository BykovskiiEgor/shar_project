from docx import Document
from pdfcreator.application.interfaces.document_reader import DocumentReader

class DocxReader(DocumentReader):
    def read(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    
    def read_tables(self, file_path: str) -> list:        
        doc = Document(file_path)
        tables = doc.tables
        
        all_tables_data = [] 
        
        for table in tables:
            table_data = []  
            
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            all_tables_data.append(table_data)
        
        return all_tables_data