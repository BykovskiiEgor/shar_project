from pdfcreator.domain.dto.document import DocumentInputDTO
from docx import Document

class DocumentService:
    def validate_format(self, document: DocumentInputDTO):
        """Check document format"""
        if not document.file_path.endswith('.docx'):
            raise ValueError("Invalid document format. Only .docx files are supported.")
       
       
    def validate_structure(self, document: DocumentInputDTO):
        """Checks that file is not empty"""
        if not self.__has_text(document) and not self.__has_tables(document):
            raise ValueError("Document is empty")
       
       
    def __has_text(self, document: DocumentInputDTO) -> bool:
        """Check if document has text"""
        doc = Document(document.file_path)
        return any(paragraph.text.strip() for paragraph in doc.paragraphs)
    
    
    def __has_tables(self, document: DocumentInputDTO) -> bool:    
        """Check if document has tables"""
        doc = Document(document.file_path)
        return len(doc.tables) > 0
